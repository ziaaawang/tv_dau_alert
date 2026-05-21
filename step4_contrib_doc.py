"""
Step 4: 归因文档（贡献pp四指标）
输入: 本周4周数据 + 阈值
输出: {week_label} TV端贡献pp变动分析.docx
"""
import os
import sys
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
from config import (
    CORE_DIMS, CORE_DIM_NAMES, CORE_DIM_FIELDS, CROSS_LEVELS,
    ALERT_METRICS, CONTRIB_METRICS,
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_TABLE, COLOR_RED, COLOR_GREEN,
    TARGET_YEAR, TARGET_WEEK,
    load_data, aggregate_weekly, compute_metrics, check_threshold,
    get_week_label, get_output_dir, load_pkl, THRESHOLD_FILE,
    fmt_pct, fmt_pp, fmt_dau_wan,
)


# ═══════════════════════════════════════════
# docx 工具（复用 step3 的模式）
# ═══════════════════════════════════════════

def set_run_font(run, size=FONT_SIZE_BODY, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=14 if level == 1 else 12 if level == 2 else 10.5, bold=True)
    return h


def add_para(doc, text, bold=False, color=None, size=FONT_SIZE_BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            text = str(val) if val is not None else '-'
            color = None
            if '触发' in text:
                color = COLOR_RED
            elif isinstance(val, (int, float)):
                color = COLOR_RED if val > 0 else COLOR_GREEN if val < 0 else None
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=FONT_SIZE_TABLE, color=color)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ═══════════════════════════════════════════
# 归因逻辑
# ═══════════════════════════════════════════

def get_topN_contrib(cur_df, dim_fields, metric, n=10):
    """获取指标 Top N 拖累/拉动"""
    if metric not in cur_df.columns:
        return [], []
    sorted_df = cur_df.dropna(subset=[metric]).sort_values(metric)
    drag = []
    for _, row in sorted_df.head(n).iterrows():
        label = ' | '.join(str(row[f]) for f in dim_fields)
        drag.append({'combo': label, 'value': row[metric], 'dau': row.get('dau', 0)})
    pull = []
    for _, row in sorted_df.tail(n).iloc[::-1].iterrows():
        label = ' | '.join(str(row[f]) for f in dim_fields)
        pull.append({'combo': label, 'value': row[metric], 'dau': row.get('dau', 0)})
    return drag, pull


def run(core_data_path):
    print("=" * 60)
    print("Step 4: 归因文档")
    print("=" * 60)

    # ── 数据加载 ──
    df = load_data(core_data_path)
    weekly_all = aggregate_weekly(df, CORE_DIM_FIELDS)
    total_weekly = weekly_all.groupby(['iso_year', 'iso_week'], as_index=False)['dau'].sum()
    total_weekly.rename(columns={'dau': 'total_dau'}, inplace=True)

    thresholds = load_pkl(THRESHOLD_FILE)
    week_label = get_week_label(TARGET_YEAR, TARGET_WEEK)
    output_dir = get_output_dir()

    # ── 整体指标 ──
    weekly_total = weekly_all.groupby(['iso_year', 'iso_week'], as_index=False)['dau'].sum()
    overall_metrics = compute_metrics(weekly_total, [], total_weekly)
    cur_overall = overall_metrics[
        (overall_metrics['iso_year'] == TARGET_YEAR) & (overall_metrics['iso_week'] == TARGET_WEEK)
    ]
    if cur_overall.empty:
        print("错误: 当前周无数据")
        return
    cur_overall = cur_overall.iloc[0]

    overall_t = thresholds.get('整体', {})

    # 整体贡献pp判定
    overall_triggers = {}
    overall_values = {}
    anomalous_metrics = []
    for m in CONTRIB_METRICS:
        val = cur_overall.get(m)
        overall_values[m] = val
        t_info = overall_t.get(m, {})
        status, _ = check_threshold(val, t_info.get('lo'), t_info.get('hi'))
        overall_triggers[m] = status
        if status != '正常':
            anomalous_metrics.append(m)

    # ── 创建文档 ──
    doc = Document()
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    title = doc.add_heading(f'{week_label} TV端贡献pp变动分析', level=0)
    for run in title.runs:
        set_run_font(run, size=16, bold=True)

    # ══════ 整体概览 ══════
    add_heading(doc, '1  整体概览', level=1)

    headers = ['指标', '数值', '阈值区间', '判定', '方向', '最大影响项']
    rows = []

    # 预计算单维度最大影响
    max_impact = {}
    for m in CONTRIB_METRICS:
        max_abs = 0
        max_label = '-'
        for dim_name in CORE_DIM_NAMES:
            dim_field = CORE_DIMS[dim_name]['field']
            weekly_dim = weekly_all.groupby(
                ['iso_year', 'iso_week', dim_field], as_index=False
            )['dau'].sum()
            metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
            cur_dim = metrics_dim[
                (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
            ]
            for _, row in cur_dim.iterrows():
                v = row.get(m)
                if v is not None and abs(v) > max_abs:
                    max_abs = abs(v)
                    max_label = f"{dim_name}-{row[dim_field]}（{v:+.2f}pp）"
        max_impact[m] = max_label

    for m in CONTRIB_METRICS:
        val = overall_values[m]
        t_info = overall_t.get(m, {})
        lo, hi = t_info.get('lo'), t_info.get('hi')
        val_s = f"{val:+.2f}pp" if val is not None else '-'
        lo_s = f"{lo:.2f}" if lo is not None else '-'
        hi_s = f"{hi:.2f}" if hi is not None else '-'

        direction = ''
        if m == 'WoW贡献pp':
            direction = '环比涨' if (val and val > 0) else '环比跌'
        elif m == 'YoY贡献pp':
            direction = '同比涨' if (val and val > 0) else '同比跌'
        elif m == 'YoY贡献pp周环比差值':
            direction = '同比加速' if (val and val > 0) else '同比减速'
        elif m == 'WoW贡献pp年同比差值':
            direction = '环比强于去年' if (val and val > 0) else '环比弱于去年'

        rows.append([m, val_s, f'[{lo_s}, {hi_s}]', overall_triggers[m], direction, max_impact[m]])

    add_table(doc, headers, rows)

    n_anomalous = len(anomalous_metrics)
    add_para(doc, f'\n{n_anomalous}/4 个指标超出阈值'
             + (f'（{"、".join(anomalous_metrics)}），变动显著，需重点关注。' if anomalous_metrics else '，均在正常范围内。'))

    if not anomalous_metrics:
        normal_metrics = [m for m in CONTRIB_METRICS if m not in anomalous_metrics]
        add_para(doc, f'正常指标：{"、".join(f"{m}（{overall_values[m]:+.2f}pp）" for m in normal_metrics)}')

    # ══════ 第2~N章：异常指标归因 ══════
    chapter = 2
    convergent_all = {}

    for metric in anomalous_metrics if anomalous_metrics else CONTRIB_METRICS[:1]:
        is_anomalous = metric in anomalous_metrics
        add_heading(doc, f'{chapter}  {metric}（{overall_values[metric]:+.2f}pp）', level=1)

        if not is_anomalous:
            add_para(doc, f'{metric} 在正常范围内，以下仅供参考。')

        # 单维度排序
        add_heading(doc, f'{chapter}.1  单维度主驱动', level=2)

        for dim_name in CORE_DIM_NAMES:
            dim_field = CORE_DIMS[dim_name]['field']
            level_t = thresholds.get(dim_name, {})
            weekly_dim = weekly_all.groupby(
                ['iso_year', 'iso_week', dim_field], as_index=False
            )['dau'].sum()
            metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
            cur_dim = metrics_dim[
                (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
            ].sort_values(metric, ascending=True if overall_values[metric] and overall_values[metric] < 0 else False)

            overall_val = overall_values[metric]
            headers_dim = [dim_name, 'DAU(万)', metric, '解释占比', '判定']
            rows_dim = []
            for _, row in cur_dim.iterrows():
                v = row.get(metric)
                gk = (row[dim_field],)
                gt = level_t.get(gk, {})
                m_info = gt.get(metric, {})
                status, _ = check_threshold(v, m_info.get('lo'), m_info.get('hi'))
                explain_pct = f"{v / overall_val * 100:.0f}%" if (overall_val and v) else '-'
                rows_dim.append([
                    row[dim_field],
                    fmt_dau_wan(row.get('dau')),
                    f"{v:+.2f}pp" if v is not None else '-',
                    explain_pct,
                    status,
                ])
            if rows_dim:
                add_table(doc, headers_dim, rows_dim)

        # 叉乘归因
        for cross_idx, (level_name, dim_names_l) in enumerate(CROSS_LEVELS):
            if len(dim_names_l) < 2:
                continue
            n_dims = len(dim_names_l)
            section_num = cross_idx  # 简化编号
            dim_fields = [CORE_DIMS[d]['field'] for d in dim_names_l]
            level_t = thresholds.get(level_name, {})

            weekly_level = weekly_all.groupby(
                ['iso_year', 'iso_week'] + dim_fields, as_index=False
            )['dau'].sum()
            metrics_level = compute_metrics(weekly_level, dim_fields, total_weekly)
            cur_level = metrics_level[
                (metrics_level['iso_year'] == TARGET_YEAR) & (metrics_level['iso_week'] == TARGET_WEEK)
            ]

            drag, pull = get_topN_contrib(cur_level, dim_fields, metric, 5)

            section_label = '两两叉乘' if n_dims == 2 else '三维叉乘'
            add_heading(doc, f'{chapter}.{n_dims}  {level_name} Top5', level=2)

            # 拖累
            if drag and drag[0]['value'] < 0:
                add_para(doc, '拖累 Top5：', bold=True)
                headers_t = ['组合', 'DAU(万)', metric]
                rows_t = [[d['combo'], fmt_dau_wan(d['dau']), f"{d['value']:+.2f}pp"] for d in drag[:5] if d['value'] < 0]
                if rows_t:
                    add_table(doc, headers_t, rows_t)
                    total_drag = sum(d['value'] for d in drag[:5] if d['value'] < 0)
                    add_para(doc, f'合计 {total_drag:+.2f}pp')

            # 拉动
            if pull and pull[0]['value'] > 0:
                add_para(doc, '拉动 Top5：', bold=True)
                headers_t = ['组合', 'DAU(万)', metric]
                rows_t = [[d['combo'], fmt_dau_wan(d['dau']), f"{d['value']:+.2f}pp"] for d in pull[:5] if d['value'] > 0]
                if rows_t:
                    add_table(doc, headers_t, rows_t)
                    total_pull = sum(d['value'] for d in pull[:5] if d['value'] > 0)
                    add_para(doc, f'合计 {total_pull:+.2f}pp')

        # 小结
        add_heading(doc, f'{chapter}.{len(CROSS_LEVELS)}  小结', level=2)

        # 收敛维度
        for dim_name in CORE_DIM_NAMES:
            dim_field = CORE_DIMS[dim_name]['field']
            weekly_dim = weekly_all.groupby(
                ['iso_year', 'iso_week', dim_field], as_index=False
            )['dau'].sum()
            metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
            cur_dim = metrics_dim[
                (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
            ]
            if overall_values[metric] and overall_values[metric] < 0:
                top = cur_dim.nsmallest(1, metric)
            else:
                top = cur_dim.nlargest(1, metric)
            if not top.empty:
                r = top.iloc[0]
                key = f"{dim_name}-{r[dim_field]}"
                convergent_all[key] = convergent_all.get(key, 0) + 1
                add_para(doc, f'  {dim_name}主驱动：{r[dim_field]}（{r[metric]:+.2f}pp）')

        chapter += 1

    # ══════ 总结章 ══════
    add_heading(doc, f'{chapter}  总结', level=1)

    # 信号矩阵
    add_heading(doc, f'{chapter}.1  异常扫描（信号矩阵）', level=2)
    headers = ['指标', '判定', '方向', '主因']
    rows = []
    for m in CONTRIB_METRICS:
        direction = ''
        val = overall_values.get(m)
        if m == 'WoW贡献pp':
            direction = '环比涨' if (val and val > 0) else '环比跌'
        elif m == 'YoY贡献pp':
            direction = '同比涨' if (val and val > 0) else '同比跌'
        elif m == 'YoY贡献pp周环比差值':
            direction = '同比加速' if (val and val > 0) else '同比减速'
        elif m == 'WoW贡献pp年同比差值':
            direction = '环比强于去年' if (val and val > 0) else '环比弱于去年'
        rows.append([m, overall_triggers[m], direction, max_impact[m]])
    add_table(doc, headers, rows)

    # 收敛分析
    add_heading(doc, f'{chapter}.2  收敛归因', level=2)
    convergent_sorted = sorted(convergent_all.items(), key=lambda x: -x[1])
    multi_convergent = [(k, v) for k, v in convergent_sorted if v >= 2]

    if multi_convergent:
        add_para(doc, '以下维度在多个异常指标中同时收敛：')
        for k, v in multi_convergent:
            add_para(doc, f'  {k} → {v}/{len(anomalous_metrics) if anomalous_metrics else 1} 个指标')
        add_para(doc, '多个异常指标指向同一人群，信号强度高，该人群是核心驱动因素。')
    else:
        add_para(doc, '各指标归因方向分散，无明显收敛。')

    # 核心结论
    add_heading(doc, f'{chapter}.3  核心结论', level=2)

    wow_val = overall_values.get('WoW贡献pp')
    yoy_val = overall_values.get('YoY贡献pp')
    add_para(doc, f'1. 大盘判定：环比 {"下降" if wow_val and wow_val < 0 else "上升"} {abs(wow_val or 0):.2f}%，'
             f'{"超出正常范围" if "WoW贡献pp" in anomalous_metrics else "在正常范围内"}；'
             f'同比 {"增长" if yoy_val and yoy_val > 0 else "下降"} {abs(yoy_val or 0):.2f}%，'
             f'{"超出正常范围" if "YoY贡献pp" in anomalous_metrics else "在正常范围内"}。', bold=True)

    for m in anomalous_metrics:
        add_para(doc, f'2. {m}归因：{m} = {overall_values[m]:+.2f}pp，超出阈值。主驱动维度：{max_impact[m]}。')

    if multi_convergent:
        groups = '、'.join(k for k, _ in multi_convergent[:4])
        add_para(doc, f'3. 关键人群：{groups}')

    # ── 保存 ──
    out_path = os.path.join(output_dir, f'{week_label} TV端贡献pp变动分析.docx')
    doc.save(out_path)
    print(f"\n已生成: {out_path}")


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='TV端DAU归因文档')
    parser.add_argument('core_data', help='核心3维数据文件 (CSV/Excel)')
    args = parser.parse_args()
    run(args.core_data)
