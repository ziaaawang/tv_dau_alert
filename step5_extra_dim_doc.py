"""
Step 5: 补充维度文档
输入: 补充维度数据（城市等级/年龄/性别/内容品类）
输出: {week_label} TV端补充维度DAU分析.docx
"""
import os
import sys
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
from config import (
    EXTRA_DIMS, ALERT_METRICS, CONTRIB_METRICS, ALL_METRICS,
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_TABLE, COLOR_RED, COLOR_GREEN,
    TARGET_YEAR, TARGET_WEEK,
    load_data, aggregate_weekly, compute_metrics, check_threshold,
    get_week_label, get_output_dir, load_pkl, THRESHOLD_FILE,
    fmt_pct, fmt_pp, fmt_dau_wan,
)


def set_run_font(run, size=FONT_SIZE_BODY, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=14 if level == 1 else 12 if level == 2 else 10.5, bold=True)
    return h


def add_para(doc, text, bold=False, color=None, size=FONT_SIZE_BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            text = str(val) if val is not None else '-'
            color = None
            if '触发' in text:
                color = COLOR_RED
            elif isinstance(val, (int, float)):
                color = COLOR_RED if val > 0 else COLOR_GREEN if val < 0 else None
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=FONT_SIZE_TABLE, color=color)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def analyze_extra_dim(doc, dim_name, dim_field, data_df, thresholds, chapter_num):
    """分析单个补充维度"""
    add_heading(doc, f'{chapter_num}  {dim_name}', level=1)

    weekly = aggregate_weekly(data_df, [dim_field])
    total_weekly = weekly.groupby(['iso_year', 'iso_week'], as_index=False)['dau'].sum()
    total_weekly.rename(columns={'dau': 'total_dau'}, inplace=True)

    metrics_df = compute_metrics(weekly, [dim_field], total_weekly)
    cur = metrics_df[
        (metrics_df['iso_year'] == TARGET_YEAR) & (metrics_df['iso_week'] == TARGET_WEEK)
    ]

    if cur.empty:
        add_para(doc, '当前周无数据。')
        return

    level_t = thresholds.get(f'补充_{dim_name}', {})

    # 预警指标表
    add_heading(doc, f'{chapter_num}.1  预警指标', level=2)

    headers = [dim_name, 'DAU(万)'] + ALERT_METRICS + [f'{m}_判定' for m in ALERT_METRICS]
    rows = []
    triggered_summary = {m: 0 for m in ALERT_METRICS}

    for _, row in cur.iterrows():
        r = [row[dim_field], fmt_dau_wan(row.get('dau'))]
        gk = (row[dim_field],)
        gt = level_t.get(gk, {})

        for m in ALERT_METRICS:
            v = row.get(m)
            r.append(f"{v:+.2f}" if v is not None else '-')

        for m in ALERT_METRICS:
            v = row.get(m)
            m_info = gt.get(m, {})
            status, _ = check_threshold(v, m_info.get('lo'), m_info.get('hi'))
            r.append(status)
            if status != '正常':
                triggered_summary[m] += 1

        rows.append(r)

    add_table(doc, headers, rows)

    # 触发汇总
    n_total = len(cur)
    summary_parts = []
    for m in ALERT_METRICS:
        n_triggered = triggered_summary[m]
        pct = n_triggered / n_total * 100 if n_total else 0
        summary_parts.append(f"{m} {n_triggered}/{n_total}（{pct:.0f}%）")
    add_para(doc, f'触发概览：{", ".join(summary_parts)}')

    # 贡献pp表
    add_heading(doc, f'{chapter_num}.2  贡献pp', level=2)

    headers = [dim_name, 'DAU(万)'] + CONTRIB_METRICS + [f'{m}_判定' for m in CONTRIB_METRICS]
    rows = []

    for _, row in cur.iterrows():
        r = [row[dim_field], fmt_dau_wan(row.get('dau'))]
        gk = (row[dim_field],)
        gt = level_t.get(gk, {})

        for m in CONTRIB_METRICS:
            v = row.get(m)
            r.append(f"{v:+.2f}pp" if v is not None else '-')

        for m in CONTRIB_METRICS:
            v = row.get(m)
            m_info = gt.get(m, {})
            status, _ = check_threshold(v, m_info.get('lo'), m_info.get('hi'))
            r.append(status)

        rows.append(r)

    add_table(doc, headers, rows)

    # WoW贡献pp 拖累/拉动
    if 'WoW贡献pp' in cur.columns:
        sorted_wow = cur.sort_values('WoW贡献pp')
        if len(sorted_wow) >= 3:
            drag = sorted_wow.head(3)
            pull = sorted_wow.tail(3).iloc[::-1]
            drag_text = '／'.join(
                f"{r[dim_field]} {r['WoW贡献pp']:+.2f}pp" for _, r in drag.iterrows()
            )
            pull_text = '／'.join(
                f"{r[dim_field]} {r['WoW贡献pp']:+.2f}pp" for _, r in pull.iterrows()
            )
            add_para(doc, f'WoW贡献pp — 拖累Top3：{drag_text}')
            add_para(doc, f'WoW贡献pp — 拉动Top3：{pull_text}')

    # 多指标触发聚焦
    add_heading(doc, f'{chapter_num}.3  多指标异常聚焦', level=2)

    multi_triggered = []
    for _, row in cur.iterrows():
        gk = (row[dim_field],)
        gt = level_t.get(gk, {})
        triggered = []
        for m in ALL_METRICS:
            v = row.get(m)
            m_info = gt.get(m, {})
            status, _ = check_threshold(v, m_info.get('lo'), m_info.get('hi'))
            if status != '正常':
                triggered.append(m)
        if len(triggered) >= 2:
            multi_triggered.append({
                'value': row[dim_field],
                'dau': row.get('dau', 0),
                'count': len(triggered),
                'metrics': triggered,
            })

    if multi_triggered:
        multi_triggered.sort(key=lambda x: -x['count'])
        for item in multi_triggered:
            add_para(doc,
                     f'  {item["value"]}（DAU {fmt_dau_wan(item["dau"])}万，{item["count"]}个指标触发）：'
                     f'{", ".join(item["metrics"])}')
    else:
        add_para(doc, '无2+指标同时触发的取值。')


def run(extra_data_paths):
    """
    extra_data_paths: dict, {'城市等级': path, '年龄': path, '性别': path, '内容品类': path}
    年龄和性别可以是同一个文件
    """
    print("=" * 60)
    print("Step 5: 补充维度文档")
    print("=" * 60)

    thresholds = load_pkl(THRESHOLD_FILE)
    week_label = get_week_label(TARGET_YEAR, TARGET_WEEK)
    output_dir = get_output_dir()

    doc = Document()
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    title = doc.add_heading(f'{week_label} TV端补充维度DAU分析', level=0)
    for run_obj in title.runs:
        set_run_font(run_obj, size=16, bold=True)

    chapter = 1
    for dim_name, dim_info in EXTRA_DIMS.items():
        if dim_name not in extra_data_paths:
            print(f"  跳过 {dim_name}：无数据文件")
            continue

        path = extra_data_paths[dim_name]
        field = dim_info['field']

        print(f"\n分析 {dim_name}（{path}）...")
        data_df = load_data(path)

        if field not in data_df.columns:
            print(f"  警告: 数据中无 {field} 列，跳过")
            continue

        analyze_extra_dim(doc, dim_name, field, data_df, thresholds, chapter)
        chapter += 1

    out_path = os.path.join(output_dir, f'{week_label} TV端补充维度DAU分析.docx')
    doc.save(out_path)
    print(f"\n已生成: {out_path}")


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='TV端补充维度分析')
    parser.add_argument('--city', help='城市等级数据文件')
    parser.add_argument('--age-sex', help='年龄性别数据文件')
    parser.add_argument('--content', help='内容品类数据文件')
    args = parser.parse_args()

    paths = {}
    if args.city:
        paths['城市等级'] = args.city
    if args.age_sex:
        paths['年龄'] = args.age_sex
        paths['性别'] = args.age_sex
    if args.content:
        paths['内容品类'] = args.content

    if not paths:
        print("请提供至少一个补充维度数据文件")
        sys.exit(1)

    run(paths)
