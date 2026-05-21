"""
Step 3: 预警文档
输入: 本周4周数据 + 阈值
输出: {week_label} TV端DAU预警分析.docx
"""
import os
import sys
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
from config import (
    CORE_DIMS, CORE_DIM_NAMES, CORE_DIM_FIELDS, CROSS_LEVELS,
    ALERT_METRICS, CONTRIB_METRICS, ALL_METRICS, SIGNAL_MATRIX,
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_TABLE, COLOR_RED, COLOR_GREEN,
    TARGET_YEAR, TARGET_WEEK,
    load_data, aggregate_weekly, compute_metrics, check_threshold,
    get_week_label, get_output_dir, load_pkl, THRESHOLD_FILE,
    fmt_pct, fmt_pp, fmt_dau_wan,
)


# ═══════════════════════════════════════════
# docx 工具
# ═══════════════════════════════════════════

def set_run_font(run, size=FONT_SIZE_BODY, bold=False, color=None):
    """设置 run 字体"""
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return h


def add_para(doc, text, bold=False, color=None, size=FONT_SIZE_BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            text = str(val) if val is not None else '-'

            color = None
            if '触发' in text:
                color = COLOR_RED
            elif isinstance(val, (int, float)):
                if val > 0:
                    color = COLOR_RED
                elif val < 0:
                    color = COLOR_GREEN

            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=FONT_SIZE_TABLE, color=color)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ═══════════════════════════════════════════
# 分析逻辑
# ═══════════════════════════════════════════

def classify_signal(overall_triggers, overall_values, single_dim_trigger_counts):
    """信号矩阵分类"""
    t = {m: (overall_triggers.get(m, '正常') != '正常') for m in ALERT_METRICS}

    # YoY%周环比方向
    yoy_wow_val = overall_values.get('YoY%周环比(pp)', 0)
    t['YoY%周环比(pp)_dir'] = 'neg' if yoy_wow_val < 0 else 'pos'

    # 检查局部异常
    overall_trigger_count = sum(1 for v in t.values() if isinstance(v, bool) and v)

    for name, check_func in SIGNAL_MATRIX:
        if name == '局部异常':
            if overall_trigger_count <= 1 and single_dim_trigger_counts > 3:
                return name
            continue
        if check_func and check_func(t):
            return name

    return '正常周'


def compute_deviation(value, lo, hi):
    """计算偏离幅度"""
    if lo is None or hi is None or value is None:
        return 0
    width = abs(hi - lo) if hi != lo else 1
    if value > hi:
        return (value - hi) / width
    elif value < lo:
        return (lo - value) / width
    return 0


def run(core_data_path):
    print("=" * 60)
    print("Step 3: 预警文档")
    print("=" * 60)

    # ── 数据加载 ──
    df = load_data(core_data_path)
    weekly_all = aggregate_weekly(df, CORE_DIM_FIELDS)
    total_weekly = weekly_all.groupby(['iso_year', 'iso_week'], as_index=False)['dau'].sum()
    total_weekly.rename(columns={'dau': 'total_dau'}, inplace=True)

    thresholds = load_pkl(THRESHOLD_FILE)
    week_label = get_week_label(TARGET_YEAR, TARGET_WEEK)
    output_dir = get_output_dir()

    # ── 整体指标 ──
    weekly_total = weekly_all.groupby(['iso_year', 'iso_week'], as_index=False)['dau'].sum()
    overall_metrics = compute_metrics(weekly_total, [], total_weekly)
    cur_overall = overall_metrics[
        (overall_metrics['iso_year'] == TARGET_YEAR) & (overall_metrics['iso_week'] == TARGET_WEEK)
    ]
    if cur_overall.empty:
        print("错误: 当前周无数据")
        return
    cur_overall = cur_overall.iloc[0]

    overall_t = thresholds.get('整体', {})

    # 整体触发判定
    overall_triggers = {}
    overall_values = {}
    for m in ALERT_METRICS:
        val = cur_overall.get(m)
        overall_values[m] = val
        t_info = overall_t.get(m, {})
        status, _ = check_threshold(val, t_info.get('lo'), t_info.get('hi'))
        overall_triggers[m] = status

    trigger_count = sum(1 for s in overall_triggers.values() if s != '正常')
    if trigger_count == 0:
        alert_level = '正常（0/4 触发）'
    elif trigger_count == 1:
        alert_level = '轻度异常（1/4 触发）'
    elif trigger_count == 2:
        alert_level = '中度异常（2/4 触发）'
    else:
        alert_level = f'显著异常（{trigger_count}/4 触发）'

    # ── 创建文档 ──
    doc = Document()
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 标题
    title = doc.add_heading(f'{week_label} TV端DAU预警分析', level=0)
    for run in title.runs:
        set_run_font(run, size=16, bold=True)

    # ══════ 第1章：预警总览 ══════
    add_heading(doc, '1  预警总览', level=1)

    cur_dau = cur_overall.get('dau', 0)
    add_para(doc, f'本周TV端DAU：{fmt_dau_wan(cur_dau)}万')
    add_para(doc, f'预警等级：{alert_level}', bold=True,
             color=COLOR_RED if trigger_count >= 2 else None)

    headers = ['指标', '本周值', '阈值区间', '状态']
    rows = []
    for m in ALERT_METRICS:
        val = overall_values[m]
        t_info = overall_t.get(m, {})
        lo, hi = t_info.get('lo'), t_info.get('hi')
        lo_s = f"{lo:.2f}" if lo is not None else '-'
        hi_s = f"{hi:.2f}" if hi is not None else '-'
        val_s = f"{val:+.2f}" if val is not None else '-'
        if 'pp' in m:
            val_s += 'pp'
        else:
            val_s += '%'
        rows.append([m, val_s, f'[{lo_s}, {hi_s}]', overall_triggers[m]])
    add_table(doc, headers, rows)

    # 信号矩阵分类
    single_dim_total_triggers = 0
    for dim_name in CORE_DIM_NAMES:
        level_t = thresholds.get(dim_name, {})
        dim_field = CORE_DIMS[dim_name]['field']
        weekly_dim = weekly_all.groupby(['iso_year', 'iso_week', dim_field], as_index=False)['dau'].sum()
        metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
        cur_dim = metrics_dim[
            (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
        ]
        for _, row in cur_dim.iterrows():
            gk = (row[dim_field],)
            gt = level_t.get(gk, {})
            for m in ALERT_METRICS:
                val = row.get(m)
                t_info = gt.get(m, {})
                status, _ = check_threshold(val, t_info.get('lo'), t_info.get('hi'))
                if status != '正常':
                    single_dim_total_triggers += 1

    signal_type = classify_signal(overall_triggers, overall_values, single_dim_total_triggers)
    add_para(doc, f'\n异常性质判断：{signal_type}', bold=True)

    # ══════ 第2章：分指标预警详情 ══════
    add_heading(doc, '2  分指标预警详情', level=1)

    for m_idx, metric in enumerate(ALERT_METRICS):
        add_heading(doc, f'2.{m_idx + 1}  {metric}', level=2)

        # 整体值
        val = overall_values[metric]
        status = overall_triggers[metric]
        val_s = f"{val:+.2f}" if val is not None else '-'
        add_para(doc, f'整体值：{val_s}，判定：{status}',
                 color=COLOR_RED if status != '正常' else None)

        # 单维度触发表
        add_heading(doc, '单维度触发汇总', level=3)
        for dim_name in CORE_DIM_NAMES:
            dim_field = CORE_DIMS[dim_name]['field']
            level_t = thresholds.get(dim_name, {})

            weekly_dim = weekly_all.groupby(['iso_year', 'iso_week', dim_field], as_index=False)['dau'].sum()
            metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
            cur_dim = metrics_dim[
                (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
            ]

            headers = [dim_name, 'DAU(万)', metric, '阈值区间', '判定']
            rows = []
            for _, row in cur_dim.iterrows():
                gk = (row[dim_field],)
                gt = level_t.get(gk, {})
                m_info = gt.get(metric, {})
                lo, hi = m_info.get('lo'), m_info.get('hi')
                mv = row.get(metric)
                status, _ = check_threshold(mv, lo, hi)
                rows.append([
                    row[dim_field],
                    fmt_dau_wan(row.get('dau')),
                    f"{mv:+.2f}" if mv is not None else '-',
                    f"[{lo:.2f}, {hi:.2f}]" if lo is not None else '-',
                    status,
                ])
            if rows:
                add_table(doc, headers, rows)

        # 叉乘触发概览
        add_heading(doc, '叉乘触发概览', level=3)

        cross_summary_headers = ['层级', '总组合数', '触发数', '触发率']
        cross_summary_rows = []
        all_triggered_combos = []

        for level_name, dim_names_l in CROSS_LEVELS:
            if len(dim_names_l) < 2:
                continue
            dim_fields = [CORE_DIMS[d]['field'] for d in dim_names_l]
            level_t = thresholds.get(level_name, {})

            weekly_level = weekly_all.groupby(
                ['iso_year', 'iso_week'] + dim_fields, as_index=False
            )['dau'].sum()
            metrics_level = compute_metrics(weekly_level, dim_fields, total_weekly)
            cur_level = metrics_level[
                (metrics_level['iso_year'] == TARGET_YEAR) & (metrics_level['iso_week'] == TARGET_WEEK)
            ]

            total_combos = len(cur_level)
            triggered = 0
            for _, row in cur_level.iterrows():
                gk = tuple(row[f] for f in dim_fields)
                gt = level_t.get(gk, {})
                m_info = gt.get(metric, {})
                mv = row.get(metric)
                status, dev = check_threshold(mv, m_info.get('lo'), m_info.get('hi'))
                if status != '正常':
                    triggered += 1
                    all_triggered_combos.append({
                        'level': level_name,
                        'combo': ' | '.join(str(row[f]) for f in dim_fields),
                        'dau': row.get('dau', 0),
                        'value': mv,
                        'deviation': dev,
                    })

            pct = triggered / total_combos * 100 if total_combos else 0
            cross_summary_rows.append([level_name, total_combos, triggered, f'{pct:.0f}%'])

        if cross_summary_rows:
            add_table(doc, cross_summary_headers, cross_summary_rows)

        # Top 5 偏离
        if all_triggered_combos:
            all_triggered_combos.sort(key=lambda x: x['deviation'], reverse=True)
            add_para(doc, f'\nTop 5 偏离最大的组合：', bold=True)
            top5_headers = ['层级', '组合', 'DAU(万)', metric, '偏离幅度']
            top5_rows = []
            for item in all_triggered_combos[:5]:
                top5_rows.append([
                    item['level'],
                    item['combo'],
                    fmt_dau_wan(item['dau']),
                    f"{item['value']:+.2f}" if item['value'] is not None else '-',
                    f"{item['deviation']:.2f}",
                ])
            add_table(doc, top5_headers, top5_rows)

    # ══════ 第3章：异常归因 ══════
    add_heading(doc, '3  异常归因（多指标交叉）', level=1)

    # 3.1 单维度聚焦
    add_heading(doc, '3.1  单维度异常聚焦', level=2)
    add_para(doc, '以下维度取值在2+个预警指标同时触发：')

    for dim_name in CORE_DIM_NAMES:
        dim_field = CORE_DIMS[dim_name]['field']
        level_t = thresholds.get(dim_name, {})
        weekly_dim = weekly_all.groupby(['iso_year', 'iso_week', dim_field], as_index=False)['dau'].sum()
        metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
        cur_dim = metrics_dim[
            (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
        ]

        for _, row in cur_dim.iterrows():
            gk = (row[dim_field],)
            gt = level_t.get(gk, {})
            triggered_metrics = []
            for m in ALERT_METRICS:
                mv = row.get(m)
                m_info = gt.get(m, {})
                status, _ = check_threshold(mv, m_info.get('lo'), m_info.get('hi'))
                if status != '正常':
                    triggered_metrics.append(m)
            if len(triggered_metrics) >= 2:
                add_para(doc,
                         f'  [{dim_name}] {row[dim_field]}（DAU {fmt_dau_wan(row.get("dau"))}万）：'
                         f'触发 {", ".join(triggered_metrics)}')

    # 3.2 叉乘异常聚焦
    add_heading(doc, '3.2  叉乘异常聚焦', level=2)

    multi_trigger_combos = []
    for level_name, dim_names_l in CROSS_LEVELS:
        if len(dim_names_l) < 2:
            continue
        dim_fields = [CORE_DIMS[d]['field'] for d in dim_names_l]
        level_t = thresholds.get(level_name, {})
        weekly_level = weekly_all.groupby(
            ['iso_year', 'iso_week'] + dim_fields, as_index=False
        )['dau'].sum()
        metrics_level = compute_metrics(weekly_level, dim_fields, total_weekly)
        cur_level = metrics_level[
            (metrics_level['iso_year'] == TARGET_YEAR) & (metrics_level['iso_week'] == TARGET_WEEK)
        ]

        for _, row in cur_level.iterrows():
            gk = tuple(row[f] for f in dim_fields)
            gt = level_t.get(gk, {})
            triggered = []
            for m in ALERT_METRICS:
                mv = row.get(m)
                m_info = gt.get(m, {})
                status, _ = check_threshold(mv, m_info.get('lo'), m_info.get('hi'))
                if status != '正常':
                    triggered.append(m)
            if len(triggered) >= 2:
                multi_trigger_combos.append({
                    'level': level_name,
                    'n_dims': len(dim_names_l),
                    'combo': ' | '.join(str(row[f]) for f in dim_fields),
                    'dau': row.get('dau', 0),
                    'trigger_count': len(triggered),
                    'triggered_metrics': triggered,
                })

    multi_trigger_combos.sort(key=lambda x: (-x['trigger_count'], x['n_dims']))

    if multi_trigger_combos:
        headers = ['层级', '组合', 'DAU(万)', '触发数', '触发指标']
        rows = []
        for item in multi_trigger_combos[:20]:
            rows.append([
                item['level'], item['combo'],
                fmt_dau_wan(item['dau']), item['trigger_count'],
                '、'.join(item['triggered_metrics']),
            ])
        add_table(doc, headers, rows)
    else:
        add_para(doc, '无2+指标同时触发的叉乘组合。')

    # 3.3 异常收敛分析
    add_heading(doc, '3.3  异常收敛分析', level=2)

    dim_value_freq = {}
    # 从单维度
    for dim_name in CORE_DIM_NAMES:
        dim_field = CORE_DIMS[dim_name]['field']
        level_t = thresholds.get(dim_name, {})
        weekly_dim = weekly_all.groupby(['iso_year', 'iso_week', dim_field], as_index=False)['dau'].sum()
        metrics_dim = compute_metrics(weekly_dim, [dim_field], total_weekly)
        cur_dim = metrics_dim[
            (metrics_dim['iso_year'] == TARGET_YEAR) & (metrics_dim['iso_week'] == TARGET_WEEK)
        ]
        for _, row in cur_dim.iterrows():
            gk = (row[dim_field],)
            gt = level_t.get(gk, {})
            for m in ALERT_METRICS:
                mv = row.get(m)
                m_info = gt.get(m, {})
                status, _ = check_threshold(mv, m_info.get('lo'), m_info.get('hi'))
                if status != '正常':
                    key = f"{dim_name}-{row[dim_field]}"
                    dim_value_freq[key] = dim_value_freq.get(key, 0) + 1

    # 从叉乘
    for item in multi_trigger_combos:
        parts = item['combo'].split(' | ')
        level_dims = [d for _, dims in CROSS_LEVELS if _ == item['level'] for d in dims]
        for i, val in enumerate(parts):
            if i < len(level_dims):
                key = f"{level_dims[i]}-{val}"
                dim_value_freq[key] = dim_value_freq.get(key, 0) + item['trigger_count']

    convergent = {k: v for k, v in dim_value_freq.items() if v >= 3}
    if convergent:
        sorted_conv = sorted(convergent.items(), key=lambda x: -x[1])
        add_para(doc, '以下维度取值在多个层级/指标中反复出现（频次>=3），为收敛维度：')
        headers = ['维度-取值', '出现频次']
        rows = [[k, v] for k, v in sorted_conv]
        add_table(doc, headers, rows)
    else:
        add_para(doc, '未发现明显收敛的维度取值。')

    # ══════ 第4章：结论与建议 ══════
    add_heading(doc, '4  结论与建议', level=1)

    add_para(doc, f'异常性质：{signal_type}', bold=True)

    if convergent:
        conv_text = '、'.join(sorted_conv[0][0] for _ in [1])
        add_para(doc, f'收敛维度：{", ".join(k for k, _ in sorted_conv[:5])}')

    add_para(doc, '\n行动建议：')
    if signal_type == '正常周':
        add_para(doc, '  - 各指标均在正常范围内，持续监控即可。')
    elif signal_type == '短期波动':
        add_para(doc, '  - 短期波动，预计后续1~2周自然回归。建议持续观察回升速度。')
    elif signal_type == '季节性偏差':
        add_para(doc, '  - 本次季节性冲击强于去年同期，关注是否有新增结构性变化。')
    elif signal_type in ('趋势性恶化', '复合异常'):
        add_para(doc, '  - 需重点关注，建议深入排查收敛维度对应的业务变化。')
        add_para(doc, '  - 若后续2周内WoW%未回归正常区间，需重新评估是否叠加了其他因素。')
    else:
        add_para(doc, '  - 建议结合业务背景进一步分析收敛维度。')

    # ── 保存 ──
    out_path = os.path.join(output_dir, f'{week_label} TV端DAU预警分析.docx')
    doc.save(out_path)
    print(f"\n已生成: {out_path}")


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='TV端DAU预警文档')
    parser.add_argument('core_data', help='核心3维数据文件 (CSV/Excel)')
    args = parser.parse_args()
    run(args.core_data)
